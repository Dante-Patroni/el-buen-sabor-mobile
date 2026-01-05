# -*- coding: utf-8 -*-
"""
Script para convertir el documento Markdown a Word
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Leer el archivo markdown
    with open('flujo_informacion_el_buen_sabor.md', 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Procesar línea por línea
    lineas = contenido.split('\n')
    
    for linea in lineas:
        linea = linea.strip()
        
        if not linea:
            continue
        
        # Título principal (# )
        if linea.startswith('# ') and not linea.startswith('## '):
            texto = linea.replace('# ', '')
            p = doc.add_heading(texto, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo nivel 2 (## )
        elif linea.startswith('## ') and not linea.startswith('### '):
            texto = linea.replace('## ', '')
            doc.add_heading(texto, level=2)
        
        # Subtítulo nivel 3 (### )
        elif linea.startswith('### ') and not linea.startswith('#### '):
            texto = linea.replace('### ', '')
            doc.add_heading(texto, level=3)
        
        # Subtítulo nivel 4 (#### )
        elif linea.startswith('#### '):
            texto = linea.replace('#### ', '')
            doc.add_heading(texto, level=4)
        
        # Separadores
        elif linea.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Bloques de código (ignorar los delimitadores)
        elif linea.startswith('```'):
            continue
        
        # Listas con viñetas
        elif linea.startswith('- ') or linea.startswith('* '):
            texto = linea[2:]
            # Limpiar markdown básico
            texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)  # Negrita
            texto = re.sub(r'`(.*?)`', r'\1', texto)  # Código inline
            doc.add_paragraph(texto, style='List Bullet')
        
        # Listas numeradas
        elif re.match(r'^\d+\.', linea):
            texto = re.sub(r'^\d+\.\s*', '', linea)
            texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
            texto = re.sub(r'`(.*?)`', r'\1', texto)
            doc.add_paragraph(texto, style='List Number')
        
        # Texto normal
        else:
            # Ignorar bloques mermaid
            if 'mermaid' in linea or 'graph' in linea or 'sequenceDiagram' in linea:
                continue
            if linea.startswith('participant ') or '-->' in linea or '->>' in linea:
                continue
            
            # Limpiar markdown
            texto = re.sub(r'\*\*(.*?)\*\*', r'\1', linea)  # Negrita
            texto = re.sub(r'`(.*?)`', r'\1', texto)  # Código inline
            texto = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', texto)  # Links
            
            if texto:
                p = doc.add_paragraph(texto)
    
    # Guardar documento
    doc.save('flujo_informacion_el_buen_sabor.docx')
    print("✅ Documento Word creado exitosamente: flujo_informacion_el_buen_sabor.docx")
    
except ImportError:
    print("❌ Error: La librería python-docx no está instalada.")
    print("Instalando python-docx...")
    import subprocess
    subprocess.run(['pip', 'install', 'python-docx'], check=True)
    print("✅ Instalación completada. Por favor ejecuta el script nuevamente.")
except Exception as e:
    print(f"❌ Error: {e}")
